"""
文档处理服务

提供文档解析、分块、向量化和存储功能
"""
import os
import json
import asyncio
import logging
import uuid
from typing import List, Dict, Any, Optional
from pathlib import Path

from langchain_text_splitters import RecursiveCharacterTextSplitter
from sqlalchemy import text as sql_text

from config import Config
from services.db_manager import db_manager
from services.embedding_service import embedding_service
from models.knowledge_base import KnowledgeBase, Document, DocumentChunk

logger = logging.getLogger(__name__)


class DocumentService:
    """文档处理服务"""

    def __init__(self):
        """初始化服务"""
        self.chunk_size = Config.CHUNK_SIZE
        self.chunk_overlap = Config.CHUNK_OVERLAP
        self.upload_folder = Config.UPLOAD_FOLDER

        # 确保上传目录存在
        os.makedirs(self.upload_folder, exist_ok=True)

        # 初始化递归分词器
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=[
                "\n\n",  # 首先尝试按段落分割
                "\n",    # 然后按行分割
                "。",    # 中文句号
                "！",    # 中文感叹号
                "？",    # 中文问号
                "；",    # 中文分号
                ".",     # 英文句号
                "!",     # 英文感叹号
                "?",     # 英文问号
                ";",     # 英文分号
                " ",     # 空格
                ""       # 最后按字符分割
            ]
        )

    def extract_text_from_file(self, file_path: str) -> Optional[str]:
        """
        从文件中提取文本内容

        Args:
            file_path: 文件路径

        Returns:
            提取的文本内容，失败返回 None
        """
        try:
            file_ext = Path(file_path).suffix.lower()

            if file_ext in ['.txt', '.md', '.json', '.csv']:
                # 纯文本文件
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()

            elif file_ext == '.pdf':
                # PDF 文件 - 需要 PyPDF2 或 pdfplumber
                try:
                    import pdfplumber
                    import logging
                    # 抑制 pdfplumber 的字体警告
                    logging.getLogger('pdfminer').setLevel(logging.ERROR)
                    with pdfplumber.open(file_path) as pdf:
                        text = ""
                        for page in pdf.pages:
                            text += page.extract_text() or ""
                        return text
                except ImportError:
                    logger.warning("pdfplumber not installed, trying PyPDF2")
                    try:
                        from PyPDF2 import PdfReader
                        reader = PdfReader(file_path)
                        text = ""
                        for page in reader.pages:
                            text += page.extract_text() or ""
                        return text
                    except ImportError:
                        logger.error("Neither pdfplumber nor PyPDF2 is installed")
                        return None

            elif file_ext in ['.doc', '.docx']:
                # Word 文档 - 需要 python-docx
                try:
                    from docx import Document as DocxDocument
                    doc = DocxDocument(file_path)
                    text = "\n".join([para.text for para in doc.paragraphs])
                    return text
                except ImportError:
                    logger.error("python-docx not installed")
                    return None

            else:
                # 尝试作为文本文件读取
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        return f.read()
                except:
                    logger.error(f"Unsupported file type: {file_ext}")
                    return None

        except Exception as e:
            logger.error(f"Error extracting text from file: {str(e)}")
            return None

    def split_text(self, text: str) -> List[str]:
        """
        使用递归分词器分割文本

        Args:
            text: 原始文本

        Returns:
            分块列表
        """
        chunks = self.text_splitter.split_text(text)
        return chunks

    async def process_document(
        self,
        file_path: str,
        filename: str,
        knowledge_base_id: str
    ) -> Optional[Dict[str, Any]]:
        """
        处理文档：提取文本、分块、向量化、存储

        Args:
            file_path: 文件路径
            filename: 原始文件名
            knowledge_base_id: 知识库ID

        Returns:
            处理结果字典
        """
        # 第一阶段：预处理（在事务外进行，避免长事务）
        # 提取文本
        extracted_text = self.extract_text_from_file(file_path)
        if not extracted_text:
            return {
                'status': 'failed',
                'error': '无法提取文本内容'
            }

        # 分块
        chunks = self.split_text(extracted_text)
        if not chunks:
            return {
                'status': 'failed',
                'error': '文本分块失败'
            }

        # 批量获取向量（可能耗时的网络操作，在事务外进行）
        embeddings = await embedding_service.get_embeddings(chunks)
        if not embeddings or len(embeddings) != len(chunks):
            return {
                'status': 'failed',
                'error': '向量化失败'
            }

        # 第二阶段：数据库操作（短事务）
        document_id = uuid.uuid4()
        
        with db_manager.get_session() as session:
            try:
                # 创建文档记录
                document = Document(
                    id=document_id,
                    knowledge_base_id=knowledge_base_id,
                    filename=filename,
                    file_path=file_path,
                    file_type=Path(filename).suffix.lower(),
                    file_size=os.path.getsize(file_path),
                    status='processing',
                    content=extracted_text,
                    chunk_count=len(chunks)
                )
                session.add(document)
                session.flush()  # 获取 document.id

                # 保存分块（使用原生 SQL 以支持 vector 类型）
                for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                    chunk_id = str(uuid.uuid4())
                    embedding_str = '[' + ','.join(map(str, embedding)) + ']'
                    metadata_json = json.dumps({
                        'filename': filename,
                        'chunk_index': i,
                        'total_chunks': len(chunks)
                    })
                    
                    # 使用字符串格式化处理 vector 类型转换，避免 SQLAlchemy 参数绑定冲突
                    # 注意：embedding_str 是我们控制的数值数组字符串，不存在注入风险
                    insert_sql = f"""
                        INSERT INTO document_chunks 
                        (id, document_id, knowledge_base_id, chunk_index, content, embedding, chunk_metadata, created_at)
                        VALUES (:id, :doc_id, :kb_id, :idx, :content, '{embedding_str}'::vector, :metadata, NOW())
                    """
                    session.execute(sql_text(insert_sql), {
                        'id': chunk_id,
                        'doc_id': str(document.id),
                        'kb_id': knowledge_base_id,
                        'idx': i,
                        'content': chunk,
                        'metadata': metadata_json
                    })

                # 更新文档状态
                document.status = 'completed'

                # 更新知识库统计（使用原生 SQL）
                session.execute(sql_text("""
                    UPDATE knowledge_bases 
                    SET document_count = document_count + 1, 
                        chunk_count = chunk_count + :chunk_count,
                        updated_at = NOW()
                    WHERE id = :kb_id
                """), {'chunk_count': len(chunks), 'kb_id': knowledge_base_id})

                return {
                    'status': 'completed',
                    'document_id': str(document.id),
                    'chunk_count': len(chunks)
                }

            except Exception as e:
                logger.error(f"Error processing document: {str(e)}")
                # 事务会自动回滚，不需要手动处理
                raise e

    def get_supported_file_types(self) -> List[str]:
        """
        获取支持的文件类型

        Returns:
            支持的文件扩展名列表
        """
        return ['.txt', '.md', '.json', '.csv', '.pdf', '.doc', '.docx']


# 创建全局实例
document_service = DocumentService()

__all__ = ['DocumentService', 'document_service']
